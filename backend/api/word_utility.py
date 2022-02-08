import datetime
import locale
import uuid

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

from organisation.models import Organisation, TypeOrganisation


def get_main_electric(document, data):
    header = [
        'на проведение приемо-сдаточных, эксплуатационных испытаний электрооборудования до 1000В, расположенного по адресу:'
    ]
    for header_str in header:
        p = document.add_paragraph()
        font = p.add_run(header_str).font
        font.bold = True
        font.size = Pt(14)
        font.name = 'Times New Roman'
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.add_run('	'+data.get('obj_address')+'	').underline = True
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tab_stops = pf.tab_stops
    tab_stop = tab_stops.add_tab_stop(Cm(9.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stop = tab_stops.add_tab_stop(Cm(18.5), WD_TAB_ALIGNMENT.CENTER)

    p = document.add_paragraph()
    p.add_run('в следующем объеме:').bold = True

    for i, control in enumerate(data.get('controls')):
        p = document.add_paragraph('{}. {}'.format(i+1, control))
    p.text = p.text[:-1] + '.'


def get_main_lift(document, data):
    header = [
        'на проведение оценки соответствия лифта требованиям технического регламента Таможенного союза «Безопасность лифтов»'
    ]
    for header_str in header:
        p = document.add_paragraph()
        font = p.add_run(header_str).font
        font.bold = True
        font.size = Pt(14)
        font.name = 'Times New Roman'
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.add_run('Просим Вас провести оценку соответствия в форме ')
    r = p.add_run(data.get('type_work'))
    r.bold = True
    r.underline = True
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = document.add_table(rows=1, cols=8, style='Table Grid')
    cells = table.rows[0].cells
    p = cells[0].paragraphs[0]
    p.add_run('№ п/п').bold
    p = cells[1].paragraphs[0]
    p.add_run('Адрес установки лифта(ов)').bold
    p = cells[2].paragraphs[0]
    p.add_run('Рег.№').bold
    p = cells[3].paragraphs[0]
    p.add_run('Год ввода').bold
    p = cells[4].paragraphs[0]
    p.add_run('Тип').bold
    p = cells[5].paragraphs[0]
    p.add_run('Грузоподъемность').bold
    p = cells[6].paragraphs[0]
    p.add_run('Кол-во остановок').bold
    p = cells[7].paragraphs[0]
    p.add_run('Месяц и год последнего освидетельствования').bold

    for i, obj in enumerate(data.get('table_rows')):
        cells = table.add_row().cells
        cells[0].text = str(i+1)
        cells[1].text = obj.get('address')
        cells[2].text = obj.get('reg_num')
        cells[3].text = obj.get('mf_year')
        cells[4].text = obj.get('type_lift')
        cells[5].text = obj.get('capacity')
        cells[6].text = obj.get('floors')
        cells[7].text = obj.get('date_exam')

    document.add_paragraph()
    document.add_paragraph()


def get_main_NK(document, data):
    header = [
        'на проведение неразрушающего контроля'
    ]
    for header_str in header:
        p = document.add_paragraph()
        font = p.add_run(header_str).font
        font.bold = True
        font.size = Pt(14)
        font.name = 'Times New Roman'
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    p = document.add_paragraph()
    p.add_run('1)Наименование оборудования (объектов)').bold = True
    for obj in data.get('objects'):
       document.add_paragraph(obj)

    p = document.add_paragraph()
    p.add_run('2)Вид (метод) неразрушающего контроля').bold = True
    for method in data.get('methods'):
        document.add_paragraph(method)

    table = document.add_table(rows=1, cols=5, style='Table Grid')
    cells = table.rows[0].cells
    p = cells[0].paragraphs[0]
    p.add_run('№ п/п').bold
    p = cells[1].paragraphs[0]
    p.add_run('Адрес нахождения объекта контроля').bold
    p = cells[2].paragraphs[0]
    p.add_run('Наименование объекта контроля').bold
    p = cells[3].paragraphs[0]
    p.add_run('Наименование элементов подвергаемых контролю').bold
    p = cells[4].paragraphs[0]
    p.add_run('Объем (количество)').bold

    for i, row in enumerate(data.get('table1_rows')):
        cells = table.add_row().cells
        cells[0].text = str(i+1)
        cells[1].text = row.get('address')
        cells[2].text = row.get('object')
        cells[3].text = row.get('element')
        cells[4].text = row.get('count')

    document.add_paragraph()


def get_request_template(data):
    document = Document()
    header = [
        '',
        '',
        'Директору',
        'ООО «ИКЦ «Запсиб-Экспертиза»',
        'А.Г. Безденежных',
        '',
        '',
    ]
    current_section = document.sections[-1]
    current_section.left_margin = Cm(2)
    current_section.right_margin = Cm(1)
    current_section.top_margin = Cm(1)
    current_section.bottom_margin = Cm(1)

    for header_str in header:
        p = document.add_paragraph()
        font = p.add_run(header_str).font
        font.size = Pt(14)
        font.name = 'Times New Roman'
        pf = p.paragraph_format
        pf.left_indent = Cm(10)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    p = document.add_paragraph()
    p.add_run('ЗАЯВКА').bold = True
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if 'неразрушающего контроля' in data.get('work'):
        get_main_NK(document, data)
    elif 'оценки соответствия лифта' in data.get('work'):
        get_main_lift(document, data)
    elif 'испытаний электрооборудования' in data.get('work'):
        get_main_electric(document, data)


    customer = [
        {
            'name': 'Наименование учреждения/организации-плательщика (в соответствии с учредительными документами):',
            'data': "{} {}".format(data['customer'].get('type_customer'), data['customer'].get('full_name'))
        },
        {
            'name': 'Должность, ФИО руководителя:',
            'data': "{}, {} {} {}".format(
                data['customer'].get('head'),
                data['customer'].get('head_lastname'),
                data['customer'].get('head_name'),
                data['customer'].get('head_surname')
            )
        },
        {
            'name': 'Основание для подписания договора руководителем или иным уполномоченным лицом: устав, приказ, доверенность (с приложением заверенной копии доверенности):',
            'data': 'Устав'
        },
        {
            'name': 'Юридический адрес:',
            'data': data['customer'].get('legal_address')
        },
        {
            'name': 'Почтовый адрес:',
            'data': data['customer'].get('post_address')
        },
        {
            'name': 'Расчетный счет:',
            'data': data['bank'].get('payment_account'),
            'inline': True
        },
        {
            'name': 'Наименование банка:',
            'data': data['bank'].get('name'),
            'inline': True
        },
        {
            'name': 'Кор. Счет:',
            'data': data['bank'].get('correspondent_account'),
            'inline': True
        },
        {
            'name': 'БИК:',
            'data': data['bank'].get('bic'),
            'inline': True
        },
        {
            'name': 'ИНН/КПП(организации):',
            'data': '{}/{}'.format(data['customer'].get('inn'), data['customer'].get('kpp')),
            'inline': True
        },
        {
            'name': 'ОГРН:',
            'data': data['customer'].get('ogrn'),
            'inline': True
        },
        {
            'name': 'Тел/факс:',
            'data': data['customer'].get('phone'),
            'inline': True
        },
        {
            'name': 'e-mail:',
            'data': data['customer'].get('email'),
            'inline': True
        },
    ]
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    k = 0
    i = 0
    cells = table.rows[0].cells
    p = cells[0].paragraphs[0]
    r = p.add_run('Реквизиты заказчика:')
    r.bold = True
    r.underline = True
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p = cells[1].paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for prop in customer:
        if prop.get('inline'):
            if i == 0:
                cells = table.add_row().cells
            p = cells[i].paragraphs[0]
            p.add_run(prop['name']).bold = True
            p.add_run(prop['data'])
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.keep_together = True
            i = i ^ 1
        else:
            cells = table.add_row().cells
            p = cells[0].paragraphs[0]
            p.add_run(prop['name']).bold = True
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            p = cells[1].paragraphs[0]
            p.add_run(prop['data'])
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.keep_together = True

    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run('		').underline = True
    p.add_run(' / ')
    p.add_run('		').underline = True
    p.add_run(' / ')
    p.add_run('		').underline = True
    pf = p.paragraph_format
    tab_stops = pf.tab_stops
    tab_stop = tab_stops.add_tab_stop(Cm(4), WD_TAB_ALIGNMENT.CENTER)
    tab_stop = tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.RIGHT)
    tab_stop = tab_stops.add_tab_stop(Cm(10), WD_TAB_ALIGNMENT.CENTER)
    tab_stop = tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.RIGHT)
    tab_stop = tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
    tab_stop = tab_stops.add_tab_stop(Cm(18), WD_TAB_ALIGNMENT.RIGHT)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p = document.add_paragraph()
    p.add_run('	(должность, наименование организации)		(подпись)		(Ф.И.О.)	').font.superscript = True
    pf = p.paragraph_format
    tab_stops = pf.tab_stops
    tab_stop = tab_stops.add_tab_stop(Cm(4), WD_TAB_ALIGNMENT.CENTER)
    tab_stop = tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.RIGHT)
    tab_stop = tab_stops.add_tab_stop(Cm(10), WD_TAB_ALIGNMENT.CENTER)
    tab_stop = tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.RIGHT)
    tab_stop = tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
    tab_stop = tab_stops.add_tab_stop(Cm(18), WD_TAB_ALIGNMENT.RIGHT)

    document.add_paragraph()
    p = document.add_paragraph("Исполнитель:")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p = document.add_paragraph(data['customer'].get('contact_name'))
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p = document.add_paragraph(data['customer'].get('phone'))
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    file_name = 'Заявка.docx'
    document.save(file_name)
    return file_name


def get_contract(lead):
    type_customer = TypeOrganisation.objects.get(name="Общество с ограниченной ответственностью")
    my_organisation = Organisation.objects.filter(type_customer=type_customer).first()
    file_name = 'api/blank lifts.docx'
    document = Document('api/blank lifts.docx')
    my_organisation
    for p in document.paragraphs:
        for r in p.runs:
            r.text = r.text.replace('{%num_contract%}', str(lead.pk))
            r.text = r.text.replace('{%date_contract%}', datetime.date.today().strftime("«%d» %B %Y"))
            r.text = r.text.replace('{%city%}', 'Новокузнецк')
            r.text = r.text.replace('{%customer_full%}', lead.customer_full_name)
            r.text = r.text.replace('{%customer_head%}', '{} {} {} {}'.format(
                lead.customer_head,
                lead.customer_name,
                lead.customer_surname,
                lead.customer_lastname
            ))
            r.text = r.text.replace('{%executor_full%}', "Общество с ограниченной ответственностью «ИКЦ «Запсиб-Экспертиза»")
            r.text = r.text.replace('{%executor_head%}',
                                    "Безденежных Алексея Геннадьевича")
    document.save('file_name.docx')
    return 'file_name.docx'


if __name__ == '__main__':
    works = [
        'проведение оценки соответствия лифта требованиям технического регламента Таможенного союза «Безопасность лифтов»',
        'проведение приемо-сдаточных, эксплуатационных испытаний электрооборудования до 1000в',
        'проведение неразрушающего контроля'
    ]
    data = {
        'customer': {
            'desc': '4217108712, г. Новокузнецк, ООО "ИКЦ "ЗАПСИБ-ЭКСПЕРТИЗА"',
            'inn': '4217108712',
            'kpp': '421701001',
            'ogrn': '1084217008061',
            'phone': '89122700658',
            'email': 'galuzik@yandex.ru',
            'full_name': 'ИНЖЕНЕРНЫЙ-КОНСУЛЬТАЦИОННЫЙ ЦЕНТР ЗАПСИБ-ЭКСПЕРТИЗА',
            'type_customer': 'ООО',
            'head': 'ДИРЕКТОР',
            'head_name': 'Алексей',
            'head_surname': 'Геннадьевич',
            'head_lastname': 'Безденежных',
            'legal_address': 'Кемеровская область - Кузбасс, г Новокузнецк, р-н Центральный, пр-кт Дружбы, д 58, пом 105',
            'Description': '4217108712, г. Новокузнецк, ООО "ИКЦ "ЗАПСИБ-ЭКСПЕРТИЗА"',
            'post_address': 'Кемеровская область - Кузбасс, г Новокузнецк, р-н Центральный, пр-кт Дружбы, д 58, пом 105',
            'contact_name': 'Евгений Иванов'
        },
        'bank': {
            'desc': 'СИБИРСКИЙ БАНК ПАО СБЕРБАНК, БИК 045004641, ИНН 7707083893',
            'name': 'СИБИРСКИЙ БАНК ПАО СБЕРБАНК',
            'bic': '045004641',
            'inn': '7707083893',
            'kpp': '540602001',
            'correspondent_account': '30101810500000000641',
            'Description': 'СИБИРСКИЙ БАНК ПАО СБЕРБАНК, БИК 045004641, ИНН 7707083893',
            'payment_account': '12345'
        },
        'type_work': 'полного технического освидетельствования',
        'work': works[2],
        'controls': [
            'Проверка соответствия смонтированной схемы электроустановки требованиям нормативно-технической документации (визуальный осмотр),',
            'Измерение сопротивления изоляции кабельных линий, электродвигателей, электрических аппаратов, вторичных цепей и электропроводок напряжением до 1000 вольт,',
            'Проверка и испытание расцепителей автоматических выключателей,'
        ],
        'objects': [
            'Системы газоснабжения (газораспределения)',
            'Оборудование нефтяной и газовой промышленности',
            'Здания и сооружения (строительные объекты)'
        ],
        'methods': [
            'Ультразвуковой (дефектоскопия, толщинометрия)',
            'Проникающими веществами (капиллярный)'
        ],
        'table1_rows': [
            {
                'address': 'Парковая 56',
                'object': 'Газопровод',
                'element': 'Сварные соединения',
                'count': '50'
            },
            {
                'address': 'Репина 7',
                'object': 'Лифт',
                'element': 'Грузоподъемный механизм',
                'count': '2'
            }
        ],
        'table_rows': [
            {
                'address': 'Ленина 34',
                'reg': '',
                'year_of_commission': None,
                'type': '',
                'capacity': '250',
                'floors': '5',
                'last_verife': '',
                'dpicker': False,
                'type_lift': 'Пассажирский',
                'reg_num': '12345',
                'mf_year': '2012',
                'date_exam': '2019-10'
            },
            {
                'address': 'Мира 43',
                'reg': '',
                'year_of_commission': None,
                'type': '',
                'capacity': '500',
                'floors': '3',
                'last_verife': '',
                'dpicker': False,
                'type_lift': 'Грузовой',
                'reg_num': '54321',
                'mf_year': '2010',
                'date_exam': '2019-12'
            }
        ],
        'obj_address': 'Победы 43'
    }

    get_request_template(data)
